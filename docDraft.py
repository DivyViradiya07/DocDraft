#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import os
import groq # type: ignore
from dotenv import load_dotenv # type: ignore
from docx import Document # type: ignore


# In[2]:


load_dotenv()
api_key = os.getenv("GROQ_API_KEY")


# In[3]:


client = groq.Groq(api_key=api_key)


# FUNCTION FOR GIVING API CALL TO GROQ

# In[4]:


def generate_legal_document(prompt):
    """Handles the API call to generate legal documents."""
    response = client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=[{"role": "system", "content": prompt}]
    )
    return response.choices[0].message.content


# PROMPTS FOR GENERATION OF CORE LEGAL DOCS

# In[5]:


def generate_writ_petition(case_number, year, petitioner, respondent, court_name, jurisdiction, legal_grounds, relief_sought, supporting_documents):
    prompt = f"""
    Generate a formal writ petition in a professional legal format with the following details:
    
    Case Number: {case_number} of {year}
    Petitioner: {petitioner}
    Respondent: {respondent}
    Court: {court_name}
    Jurisdiction: {jurisdiction}
    Legal Grounds: {legal_grounds}
    Relief Sought: {relief_sought}
    Supporting Documents: {', '.join(supporting_documents)}

    Structure the writ petition with these sections:
    - Title & Case Details
    - Introduction & Parties Involved
    - Jurisdiction Basis
    - Facts of the Case
    - Legal Grounds & Precedents
    - Relief Sought
    - List of Supporting Documents
    - Prayer (Final Request to the Court)

    The document should be in a professional and structured legal format, using numbered paragraphs where necessary.
    """

    return generate_legal_document(prompt)


# In[6]:


def generate_affidavit(case_number, year, petitioner, respondent, affiant_name, order_date, property_name, statement_of_facts):
    prompt = f"""
    Generate a formal legal affidavit in a professional tone with the following details:
    
    Case Number: {case_number} of {year}
    Petitioner: {petitioner}
    Respondent: {respondent}
    Affiant: {affiant_name}
    Order Date: {order_date}
    Property Name: {property_name}
    Statement of Facts: {statement_of_facts}

    Structure the affidavit in proper legal format with these sections:
    - Title & Case Details
    - Introduction (Details of the Affiant)
    - Statement of Facts
    - Benefits Offered (if applicable, such as rent & corpus details)
    - Compliance with Regulations
    - Oath & Affirmation
    - Signature & Notarization

    The document should be structured with numbered clauses where necessary, ensuring clarity and legal precision.
    """

    return generate_legal_document(prompt)


# In[7]:


def generate_patent_application(application_number, year, inventor_name, assignee, title, field_of_invention, background, summary, claims, drawings_description):
    prompt = f"""
    Generate a formal patent application in a professional and structured legal format with the following details:

    Application Number: {application_number} of {year}
    Inventor: {inventor_name}
    Assignee: {assignee}
    Title of the Invention: {title}
    Field of Invention: {field_of_invention}
    Background: {background}
    Summary of Invention: {summary}
    Claims: {claims}
    Description of Drawings: {drawings_description}

    Structure the patent application with these sections:
    - Title & Application Details
    - Field of the Invention
    - Background of the Invention
    - Summary of the Invention
    - Detailed Description of the Invention
    - Claims (clearly defining the scope of protection)
    - Brief Description of Drawings (if applicable)
    - Abstract

    Ensure that the claims are structured clearly and precisely, with numbered points outlining each aspect of the invention.
    """

    return generate_legal_document(prompt)


# In[8]:


def generate_annexure(case_number, year, petitioner, respondent, annexure_title, annexure_number, description, supporting_documents):
    prompt = f"""
    Generate a formal legal annexure in a professional and structured format with the following details:
    
    Case Number: {case_number} of {year}
    Petitioner: {petitioner}
    Respondent: {respondent}
    Annexure Title: {annexure_title}
    Annexure Number: {annexure_number}
    Description: {description}
    Supporting Documents: {', '.join(supporting_documents)}

    Structure the annexure document as follows:
    - Title (Annexure Number and Case Details)
    - Introduction (Brief description of the annexure and its relevance)
    - Detailed Description (Explaining the contents of the annexure)
    - List of Supporting Documents
    - Affirmation (Declaring the authenticity of the annexure)
    - Date & Signature of the Affiant or Legal Representative

    Ensure the document is formatted in a structured legal manner, maintaining clarity and professionalism.
    """

    return generate_legal_document(prompt)



# PROMPTS FOR GENERATION OF Evidence and Compliance DOCS
# 

# In[9]:


def generate_witness_statement(case_number, year, court_name, witness_name, witness_details, statement):
    prompt = f"""
    Generate a formal witness statement in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Court: {court_name}
    Witness Name: {witness_name}
    Witness Details: {witness_details}
    Statement: {statement}

    Structure the witness statement as follows:
    - Title & Case Details
    - Witness Introduction (Name, Age, Occupation, Relationship to the Case)
    - Statement of Facts (Clear and chronological details of what the witness knows)
    - Affirmation (Witness swears under oath that the statement is true)
    - Signature & Date

    Ensure the document is formal and legally admissible.
    """

    return generate_legal_document(prompt)



# In[10]:


def generate_exhibit(case_number, year, exhibit_number, exhibit_title, description, attached_documents):
    prompt = f"""
    Generate a formal exhibit document in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Exhibit Number: {exhibit_number}
    Exhibit Title: {exhibit_title}
    Description: {description}
    Attached Documents: {', '.join(attached_documents)}

    Structure the exhibit document as follows:
    - Title (Exhibit Number and Case Details)
    - Description of Exhibit (Purpose and relevance)
    - List of Attached Documents
    - Certification of Authenticity
    - Date & Signature of the Submitting Party

    Ensure proper legal formatting for use in court.
    """

    return generate_legal_document(prompt)


# In[11]:


def generate_forensic_report(case_number, year, forensic_expert, forensic_field, report_summary, findings, conclusion):
    prompt = f"""
    Generate a formal forensic report in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Forensic Expert: {forensic_expert}
    Field of Analysis: {forensic_field}
    Report Summary: {report_summary}
    Findings: {findings}
    Conclusion: {conclusion}

    Structure the forensic report as follows:
    - Title & Case Details
    - Expert Introduction (Credentials and area of expertise)
    - Summary of the Investigation
    - Findings (Technical details, data analysis, results)
    - Conclusion (Expert opinion based on evidence)
    - Signature & Certification by the Forensic Expert

    Ensure technical accuracy and legal validity.
    """

    return generate_legal_document(prompt)


# In[12]:


def generate_expert_opinion(case_number, year, expert_name, field_of_expertise, opinion_summary, detailed_opinion, supporting_references):
    prompt = f"""
    Generate a formal expert opinion document in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Expert Name: {expert_name}
    Field of Expertise: {field_of_expertise}
    Opinion Summary: {opinion_summary}
    Detailed Opinion: {detailed_opinion}
    Supporting References: {', '.join(supporting_references)}

    Structure the expert opinion document as follows:
    - Title & Case Details
    - Expert Introduction (Qualifications & Expertise)
    - Summary of Opinion
    - Detailed Explanation with Analysis
    - Supporting References (Case Laws, Precedents, Technical Reports)
    - Conclusion
    - Signature & Date

    Ensure that the document follows professional legal standards.
    """

    return generate_legal_document(prompt)


# In[ ]:


from docx import Document

def save_to_docx(content, filename):
    """Save formatted legal document content to a .docx file."""
    doc = Document()

    lines = content.split("\n")  # Split content into lines

    for line in lines:
        line = line.strip()
        if not line:
            continue  # Skip empty lines
        
        # Check for headings (marked with ###)
        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=2)  # Level 2 Heading
        elif line.startswith("**") and line.endswith("**"):  # Bold headings
            doc.add_heading(line.replace("**", ""), level=3)  # Level 3 Heading
        else:
            doc.add_paragraph(line)  # Regular paragraph

    doc.save(filename)
    print(f"Document saved as {filename}")



# Using Jupyternotebook Widgets
# 

# In[14]:


def main():
    options = {
        "1": "Generate Writ Petition",
        "2": "Generate Affidavit",
        "3": "Generate Patent Application",
        "4": "Generate Annexure",
        "5": "Generate Witness Statement",
        "6": "Generate Exhibit",
        "7": "Generate Forensic Report",
        "8": "Generate Expert Opinion"
    }
    
    print("Select an option:")
    for key, value in options.items():
        print(f"{key}. {value}")
    
    choice = input("Enter your choice: ")
    
    if choice == "1":
        print("You are Generating Writ Petition.")
        case_number = input("Case Number: ")
        year = input("Year: ")
        petitioner = input("Petitioner: ")
        respondent = input("Respondent: ")
        court_name = input("Court Name: ")
        jurisdiction = input("Jurisdiction: ")
        legal_grounds = input("Legal Grounds: ")
        relief_sought = input("Relief Sought: ")
        supporting_documents = input("Supporting Documents (comma-separated): ").split(",")
        document = generate_writ_petition(case_number, year, petitioner, respondent, court_name, jurisdiction, legal_grounds, relief_sought, supporting_documents)
    
    elif choice == "2":
        print("You are Generating Affidavit.")
        affiant_name = input("Affiant Name: ")
        statement_of_facts = input("Statement of Facts: ")
        sworn_before = input("Sworn Before: ")
        date_of_execution = input("Date of Execution: ")
        document = generate_affidavit(affiant_name, statement_of_facts, sworn_before, date_of_execution)
        
    elif choice == "3":
        print("You are Generating Patent Application.")
        application_number = input("Application Number: ")
        year = input("Year: ")
        inventor_name = input("Inventor Name: ")
        assignee = input("Assignee: ")
        title = input("Title: ")
        field_of_invention = input("Field of Invention: ")
        background = input("Background: ")
        summary = input("Summary: ")
        claims= input("Claims: ")
        drawings_description = input("Drawings Description: ")
        document = generate_patent_application(application_number, year, inventor_name, assignee, title, field_of_invention, background, summary, claims, drawings_description)
                
    elif choice == "4":
        print("You are Generating a Legal Annexure.")
        case_number = input("Case Number: ")
        year = input("Year: ")
        petitioner = input("Petitioner: ")
        respondent = input("Respondent: ")
        annexure_title = input("Annexure Title: ")
        annexure_number = input("Annexure Number: ")
        description = input("Description: ")
        supporting_documents = input("Supporting Documents (comma-separated): ").split(",")
        document = generate_annexure(case_number, year, petitioner, respondent, annexure_title, annexure_number, description, supporting_documents)


    elif choice == "5":
        print("You are Generating a Witness Statement.")
        case_number = input("Case Number: ")
        year = input("Year: ")
        court_name = input("Court Name: ")
        witness_name = input("Witness Name: ")
        witness_details = input("Witness Details: ")
        statement = input("Statement: ")
        document = generate_witness_statement(case_number, year, court_name, witness_name, witness_details, statement)


    elif choice == "6":
        print("You are Generating an Exhibit Document.")
        case_number = input("Case Number: ")
        year = input("Year: ")
        exhibit_number = input("Exhibit Number: ")
        exhibit_title = input("Exhibit Title: ")
        description = input("Description: ")
        attached_documents = input("Attached Documents (comma-separated): ").split(",")
        document = generate_exhibit(case_number, year, exhibit_number, exhibit_title, description, attached_documents)
    
    elif choice == "7":
        print("You are Generating a Forensic Report.")
        case_number = input("Case Number: ")
        year = input("Year: ")
        forensic_expert = input("Forensic Expert: ")
        forensic_field = input("Field of Analysis: ")
        report_summary = input("Report Summary: ")
        findings = input("Findings: ")
        conclusion = input("Conclusion: ")  
        document = generate_forensic_report(case_number, year, forensic_expert, forensic_field, report_summary, findings, conclusion)

        
        
    elif choice == "8":
        print("You are Generating an Expert Opinion Document.")
        case_number = input("Case Number: ")
        year = input("Year: ")
        expert_name = input("Expert Name: ")
        field_of_expertise = input("Field of Expertise: ")
        opinion_summary = input("Opinion Summary: ")
        detailed_opinion = input("Detailed Opinion: ")
        supporting_references = input("Supporting References (comma-separated): ").split(",")
        document = generate_expert_opinion(case_number, year, expert_name, field_of_expertise, opinion_summary, detailed_opinion, supporting_references)

        
        
    else:
        print("Invalid choice!")
        return
    
    filename = input("Enter filename to save (.docx): ")
    save_to_docx(document, filename)
    
if __name__ == "__main__":
    main()


# In[ ]:




